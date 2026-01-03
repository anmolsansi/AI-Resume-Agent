import io
from pathlib import Path
from docx import Document

BASE_DIR = Path(__file__).resolve().parent.parent
OUTPUT_DIR = BASE_DIR / "output"
MASTER_RESUME_PATH = BASE_DIR / "app" / "master_resume.docx"


def _safe_company(company_name: str) -> str:
    return company_name.strip().replace(" ", "_")


def create_resume_docx(company_name: str, resume_text: str, version: int) -> Path:
    """
    Creates a versioned resume file:
      Anmol_Sansi_<Company>_v<version>.docx
    """
    try:
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

        safe_company = _safe_company(company_name)
        filename = f"Anmol_Sansi_{safe_company}_v{version}.docx"
        docx_path = OUTPUT_DIR / filename

        doc = Document()
        for line in str(resume_text).split("\n"):
            doc.add_paragraph(line)
        doc.save(docx_path)

        return docx_path
    except Exception as e:
        print(e)



def extract_text_from_docx_bytes(docx_bytes: bytes) -> str:
    doc = Document(io.BytesIO(docx_bytes))
    lines = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)
    return "\n".join(lines)


def load_master_resume_text() -> str:
    """
    Loads the master resume from app/master_resume.docx and returns plain text.
    Raises FileNotFoundError if the file doesn't exist.
    """
    if not MASTER_RESUME_PATH.exists():
        raise FileNotFoundError(
            f"Master resume not found at {MASTER_RESUME_PATH}. "
            "Create app/master_resume.docx"
        )
    doc = Document(str(MASTER_RESUME_PATH))
    lines = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)
    return "\n".join(lines)
